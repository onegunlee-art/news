# -*- coding: utf-8 -*-
"""Generate 직무 발명 설명서 Word document (Judgment 기반 시스템)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_cell_shading(cell, color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    # ASCII 파일명: 경로·도구 호환. 한글명 복제는 동일 디렉터리에 선택 저장.
    out_path = root / "docs" / "invention_disclosure_Judgment_RAG_AI.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("직무 발명 설명서 — Judgment 기반 인간 편집 재현형 AI 시스템")
    r.bold = True
    r.font.size = Pt(14)
    title.paragraph_format.space_after = Pt(18)

    # --- Table matching form layout ---
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Table Grid"

    # Row 0: (1) | (2)
    c00 = tbl.cell(0, 0)
    c01 = tbl.cell(0, 1)
    set_cell_shading(c00, "F2F2F2")
    set_cell_shading(c01, "F2F2F2")

    h1 = c00.paragraphs[0]
    r1 = h1.add_run("(1) 종래기술 도면")
    r1.bold = True
    c00.add_paragraph("도면부호: 도 1 — 종래의 즉시 검색·생성 방식 개념도")
    mermaid1 = """flowchart LR
  subgraph PRIOR["종래기술 예"]
    U["사용자 입력"] --> E["임베딩"]
    E --> V["단일 또는 단순 지식베이스 검색"]
    V --> P["프롬프트 결합"]
    P --> G["생성 모델"]
    G --> O["출력"]
  end"""
    for line in mermaid1.split("\n"):
        pp = c00.add_paragraph()
        run = pp.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        run.font.size = Pt(8)
    c00.add_paragraph("[설명] 입력 선행 판단 없이 검색·생성이 연속 수행되는 구조.")
    c00.add_paragraph("※ Mermaid 도면은 https://mermaid.live 에서 PNG/SVG로 내보내 본 칸에 삽입 가능.")

    h2 = c01.paragraphs[0]
    r2 = h2.add_run("(2) 본 발명 대표 도면")
    r2.bold = True
    c01.add_paragraph("도면부호: 도 2 — Judgment 및 Multi-Index RAG 시스템 블록도")
    mermaid2 = """flowchart TB
  Q["입력 텍스트"] --> J1["Judgment: 선행 평가"]
  J1 -->|기준 미달| CL["Clarification"]
  J1 -->|적합| EM["임베딩"]
  EM --> I1["크리틱 인덱스"]
  EM --> I2["과거 분석 인덱스"]
  EM --> I3["지식 라이브러리 인덱스"]
  EM --> I4["편집 패턴 인덱스(선택)"]
  I1 & I2 & I3 & I4 --> CB["카테고리별 컨텍스트"]
  CB --> GEN["해석·생성"]
  A1["검증"] --> A2["분석(JSON)"] --> A3["내레이션"] --> A4["편집"]
  M1["판단 패턴 저장"] --> M2["임베딩·재검색"] --> I1"""
    for line in mermaid2.split("\n"):
        pp = c01.add_paragraph()
        run = pp.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        run.font.size = Pt(7.5)
    c01.add_paragraph("[설명] 적합 시에만 멀티 인덱스 검색, 파이프라인·벡터 메모리 축적.")
    c01.add_paragraph("※ 상세 다이어그램은 mermaid.live에서 도 2로 내보내 삽입 권장.")

    # Row 1: (3) | (4)
    c10 = tbl.cell(1, 0)
    c11 = tbl.cell(1, 1)
    p3h = c10.paragraphs[0]
    p3h.add_run("(3) 종래기술의 문제점(해결해야 할 과제)").bold = True
    text3 = """일반적인 생성형 AI 및 검색 증강 생성(RAG) 시스템은 사용자 입력에 대한 목적 적합성, 도메인 적합성, 구조 완성도, 해석 가능성 등을 선행 판단하지 않은 상태에서 즉시 벡터 검색 및 생성 단계를 수행한다. 이에 따라 뉴스·시사·전략·법률·리서치 등 전문 도메인에서 부적절하거나 불명확한 입력에도 불필요한 검색 연산과 생성 연산이 반복 수행되고, 관련성이 낮은 문맥이 생성 모델에 주입되어 결과 품질 저하 및 편집 재작업 증가 문제가 발생한다.

또한 종래 기술은 단일 지식베이스 또는 단일 벡터 인덱스 중심 구조가 일반적이므로, 인간 편집자의 수정 피드백(크리틱), 과거 분석 결과, 도메인 지식 라이브러리 및 조직 내부 해석 논리 등 서로 다른 성격의 데이터를 분리·검색·융합하기 어렵다. 그 결과 조직 고유의 문체, 분석 관점, 판단 기준 및 사고 체계가 생성 결과에 충분히 반영되지 못하는 문제가 있다.

아울러 기존 생성형 AI는 대부분 단일 프롬프트 기반 생성 구조에 의존하므로, 인간 전문가가 수행하는 분석 → 해석 → 내레이션 → 편집 → 교정의 단계적 사고 및 판단 과정을 구조적으로 재현하지 못한다. 따라서 생성 과정의 중간 상태를 검증하기 어렵고, 오류 전파 및 품질 추적이 제한된다.

특히 인간 편집자의 수정·교정·비평·해석 행위가 단순 결과물 수준으로만 소비되고, 시스템 내부에서 구조화·임베딩·재검색·재활용되지 않으므로, 조직 내부의 전문 편집 역량과 사고 체계가 디지털 자산 형태로 축적되지 못하는 문제가 있다."""
    for para in text3.split("\n\n"):
        c10.add_paragraph(para.strip())

    p4h = c11.paragraphs[0]
    p4h.add_run("(4) 본 발명에서의 해결방안").bold = True
    text4 = """본 발명은 생성형 AI 시스템의 검색 증강 단계(RAG) 이전에 입력 판단(Judgment) 계층을 배치하여, 입력의 도메인 적합성, 구조 완성도, 해석 가능성, 정책 위험도 및 의미 명확성을 선행 평가한다. 판단 결과가 기준 미달인 경우에는 후속 벡터 검색 및 생성 단계를 수행하지 않고 명확화 요청(Clarification) 단계로 분기한다.

적합 판정을 받은 경우에만 복수의 이종 벡터 인덱스에 대한 조건부 검색을 수행한다. 상기 벡터 인덱스는 적어도 편집 크리틱 인덱스, 과거 분석 인덱스, 조직 지식 라이브러리 인덱스 및 인간 편집 패턴 인덱스를 포함할 수 있다.

검색 결과는 카테고리별 컨텍스트 블록으로 재구성되어 생성 모델의 프롬프트에 통합 제공되며, 이를 통해 인간 편집자의 판단 기준, 조직의 분석 논리 및 도메인별 해석 구조가 동시에 반영된다.

또한 본 발명은 검증 에이전트, 분석 에이전트, 내레이션 에이전트 및 편집 에이전트로 구성된 다단계 상태 전이 파이프라인을 통해 인간 편집 워크플로를 구조적으로 재현한다. 각 단계는 JSON 기반 구조화 스키마를 통해 상태를 전달하며, 단재별 오류 검증 및 품질 추적이 가능하도록 구성된다.

아울러 인간 편집자의 수정·교정·비평·문체 변경·논리 보정 행위를 "판단 패턴 데이터(Judgment Pattern Data)"로 구조화하여 벡터 메모리 형태로 저장하고, 이후 유사 사례 생성 시 재검색하여 생성 결과에 반영함으로써 인간 전문가의 사고 체계를 지속적으로 학습·재현하도록 구성된다."""
    # Fix typo: user had "단재별" should be "단계별" - user's original said "단계별"
    text4 = text4.replace("단재별", "단계별")
    for para in text4.split("\n\n"):
        c11.add_paragraph(para.strip())

    # Row 2: merge (5) full width
    c20 = tbl.cell(2, 0)
    c21 = tbl.cell(2, 1)
    c20.merge(c21)
    p5h = c20.paragraphs[0]
    p5h.add_run("(5) 발명의 목적").bold = True
    text5 = """본 발명의 목적은 인간 편집자 및 분석가가 수행하는 판단, 해석, 문체 교정 및 편집 의사결정 과정을 생성형 AI 시스템 내부에서 디지털 방식으로 재현·학습·재사용 가능하게 하는 데 있다.

특히 본 발명은 인간 편집자의 수정 논리, 비평 기준, 문체 교정 패턴 및 분석 구조를 단순 결과물이 아닌 "판단 데이터(Judgment Data)"로 구조화하고, 이를 벡터 메모리 형태로 축적·검색·재적용함으로써 조직 내부의 편집 역량과 사고 체계를 AI 시스템이 지속적으로 학습하도록 하는 것을 목적으로 한다.

또한 생성 이전 단계에서 입력 적합성을 판단하는 선행 Judgment 계층을 통해 인간 편집자와 유사한 검토 프로세스를 구현하고, 분석 → 내레이션 → 편집의 다단계 생성 구조를 통해 인간 수준의 콘텐츠 생성 워크플로를 재현하는 데 목적이 있다.

아울러 본 발명은 조직 단위의 편집 스타일, 분석 관점, 해석 논리 및 전략적 판단 기준을 AI가 누적 학습·모방·재생성 가능하게 함으로써, 인간 전문가 집단의 말투·표현 방식·논리 전개 및 사고 흐름을 디지털 자산화하고 재현 가능한 인공지능 기반 편집 시스템을 구현하는 것을 목적으로 한다."""
    for para in text5.split("\n\n"):
        c20.add_paragraph(para.strip())

    # Row 3: (6) full width
    c30 = tbl.cell(3, 0)
    c31 = tbl.cell(3, 1)
    c30.merge(c31)
    p6h = c30.paragraphs[0]
    p6h.add_run("(6) 발명의 구성").bold = True
    text6 = """【시스템 개요】
본 발명은 입력 판단 모듈(Judgment Module), 다중 소스 검색 모듈(Multi-Index RAG Module), 해석 생성 모듈, 다단계 상태 전이 파이프라인 모듈 및 편집 피드백 축적 모듈을 포함하는 생성형 AI 기반 인간 편집 재현 시스템에 관한 것이다.

【입력 판단 모듈(Judgment Module)】
입력 텍스트에 대해 길이·형식·도메인 적합성·구조적 완성도·정책 위험도·의미 명확성 등을 평가하고, 후속 검색 수행 여부를 동적으로 제어한다. 판단 결과는 점수(score) 또는 등급 형태로 산출될 수 있으며, 기준 미달 시 명확화 요청을 반환하고 후속 검색 및 생성 단계를 제한 또는 중단한다.

상기 판단 결과에 따라:
• 검색 수행 여부,
• 검색 깊이(depth),
• 참조 인덱스 종류,
• 생성 허용 범위,
• 후속 에이전트 활성화 여부
중 적어도 하나 이상이 동적으로 제어될 수 있다.

【다중 소스 검색 모듈(Multi-Index RAG Module)】
입력 텍스트에 대한 임베딩 벡터를 생성하고, 복수의 벡터 인덱스에 대해 독립 검색을 수행한다.

상기 복수 인덱스는:
• 편집 크리틱 인덱스,
• 과거 분석 인덱스,
• 도메인 지식 라이브러리 인덱스,
• 인간 편집 패턴 인덱스
중 적어도 하나 이상을 포함할 수 있다.

각 인덱스별 검색 결과는 카테고리 단위 컨텍스트 블록으로 재구성되며, 생성 모델의 프롬프트에 통합 제공된다.

【인간 판단 학습 구조】
본 발명은 인간 편집자의 수정·교정·비평·논리 수정·문체 변경 행위를 "판단 패턴 데이터"로 구조화하여 저장한다.

상기 판단 패턴 데이터는:
• 수정 전/후 텍스트,
• 편집 사유,
• 문체 변경 기준,
• 논리 구조 수정 정보,
• 도메인별 해석 기준,
• 조직별 편집 스타일
중 적어도 하나 이상을 포함할 수 있다.

저장된 판단 패턴 데이터는 임베딩 벡터 형태로 변환되어 벡터 메모리에 저장되며, 이후 유사 입력 발생 시 재검색되어 생성 모델의 판단 및 생성 과정에 반영된다.

이를 통해 AI 시스템은 인간 전문가의 반복적인 수정 습관과 사고 구조를 누적 학습하며, 특정 조직 또는 전문가 집단의 편집 스타일 및 말투를 점진적으로 재현하도록 구성된다.

【다단계 상태 전이 파이프라인】
검증 에이전트는 URL 및 콘텐츠 접근 가능 여부를 검증하고 본문을 추출한다.
분석 에이전트는 구조화된 분석 데이터(JSON 스키마)를 생성한다.
내레이션 에이전트는 분석 결과를 기반으로 해설용 텍스트를 생성한다.
편집 에이전트는 조직 스타일 가이드 기반 문체·톤 보정 및 논리 재구성을 수행한다.

각 단계는 상태 객체(state object)를 기반으로 순차 실행되며, 단계별 성공·실패 상태에 따라 파이프라인 분기 제어가 가능하다.

【조직 편집 메모리 시스템】
편집 피드백 및 분석 결과는 조직 단위 벡터 메모리로 축적되며, AI 시스템은 해당 메모리를 반복적으로 재검색함으로써 조직 특유의 문체·논리·판단 기준 및 인간 전문가의 말투를 지속적으로 학습한다.

이에 따라 시스템은 단순 범용 생성형 AI가 아니라 특정 조직의 인간 전문가 집단의 사고 체계 및 표현 방식을 재현 가능한 디지털 편집 에이전트로 동작할 수 있다.

【데이터 저장 구조】
벡터 데이터는 pgvector 기반 관계형 데이터베이스 또는 동등한 벡터 저장소에 저장될 수 있으며, 벡터 유사도 검색 함수에 의해 검색된다."""
    for block in text6.split("\n\n"):
        c30.add_paragraph(block.strip())

    # Row 4: (7) full width
    c40 = tbl.cell(4, 0)
    c41 = tbl.cell(4, 1)
    c40.merge(c41)
    p7h = c40.paragraphs[0]
    p7h.add_run("(7) 발명의 효과(특징)").bold = True
    text7 = """본 발명은 생성 이전 단계에서 입력 적합성을 판단함으로써 불필요한 벡터 검색 호출 및 생성 연산을 감소시키고, 문맥 오염 및 부적절한 생성 결과를 줄일 수 있다.

또한 이종 멀티 인덱스 검색 구조를 통해 편집 기준, 조직 경험, 도메인 지식 및 인간 전문가의 수정 논리를 동시에 반영할 수 있으며, 단일 지식베이스 기반 RAG 대비 문맥 정확도와 분석 일관성을 향상시킬 수 있다.

아울러 분석 → 내레이션 → 편집의 단계적 생성 구조를 통해 중간 산출물 검증 및 오류 추적이 가능하며, 생성 품질 관리 효율을 향상시킬 수 있다.

특히 인간 편집자의 수정·교정·비평·판단 과정을 구조화 데이터 및 벡터 메모리 형태로 축적함으로써, AI 시스템이 인간 전문가의 편집 논리·사고 패턴·문체·말투·표현 방식 및 전략적 해석 구조를 반복적으로 학습·재현할 수 있는 효과가 있다.

또한 조직 내부의 편집 역량과 분석 노하우가 시스템 내부 벡터 메모리로 축적되므로, 숙련된 인간 전문가의 사고 체계와 표현 스타일을 장기적으로 유지·확장·자동화할 수 있다.

이에 따라 본 발명은 단순 텍스트 생성 시스템이 아니라, 인간 전문가 집단의 판단 구조와 말투를 디지털 방식으로 재현·복제·확장 가능한 인공지능 기반 판단 및 편집 시스템을 구현할 수 있는 효과가 있다."""
    for para in text7.split("\n\n"):
        c40.add_paragraph(para.strip())

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.add_run(
        "본 문서는 초안이며, 변리사·지식재산 담당 검토 후 수정·도면 교체를 권장합니다."
    ).italic = True
    foot.runs[0].font.size = Pt(9)
    foot.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.save(out_path)
    alt = root / "docs" / "직무_발명_설명서_Judgment_기반_인간_편집_재현형_AI.docx"
    try:
        import shutil
        shutil.copy2(out_path, alt)
    except OSError:
        pass
    print(out_path)


if __name__ == "__main__":
    main()
